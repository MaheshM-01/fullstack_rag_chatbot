

import os
from pathlib import Path
from typing import List
from dataclasses import dataclass, field

# PDF reading — reads PDF pages as text
import pypdf

# Word document reading — reads .docx paragraphs
from docx import Document as DocxDocument

# HTTP client — fetches web pages
import httpx

# HTML parser — extracts clean text from HTML
from bs4 import BeautifulSoup

# Our settings object
from src.config import settings


# ============================================================
# DATA CLASS — Standardized output format
# ============================================================
@dataclass
class LoadedDocument:
    
    content: str
    metadata: dict
    source: str
    doc_type: str


# ============================================================
# MAIN LOADER CLASS
# ============================================================
class DocumentLoader:
    

    def load(self, source: str) -> List[LoadedDocument]:
        
        source = source.strip()

        # ---- URL Detection ----
       
        if source.startswith("http://") or source.startswith("https://"):
            return self._load_url(source)

        # ---- File Detection ----
        path = Path(source)

        if not path.exists():
            raise FileNotFoundError(
                f" File not found: {source}\n"
                f"   Check if path is correct."
            )

        # Get file extension (.pdf, .docx, .txt)
        extension = path.suffix.lower()

        if extension == ".pdf":
            return self._load_pdf(source)
        elif extension in [".docx", ".doc"]:
            return self._load_docx(source)
        elif extension == ".txt":
            return self._load_txt(source)
        else:
            raise ValueError(
                f" Unsupported file type: {extension}\n"
                f"   Supported types: .pdf, .docx, .doc, .txt, URLs"
            )

    # --------------------------------------------------------
    # PDF LOADER
    # --------------------------------------------------------
    def _load_pdf(self, file_path: str) -> List[LoadedDocument]:
        
        documents = []

        with open(file_path, "rb") as file:   # rb = read binary (PDF is binary)
            pdf_reader = pypdf.PdfReader(file)
            total_pages = len(pdf_reader.pages)

            print(f" Loading PDF: {Path(file_path).name} ({total_pages} pages)")

            for page_num, page in enumerate(pdf_reader.pages):

                # Extract text from this page
                text = page.extract_text()

                # Skip completely empty pages
                # (some PDFs have blank pages or image-only pages)
                if not text or text.strip() == "":
                    print(f"   ⚠️  Page {page_num + 1} is empty, skipping...")
                    continue

                documents.append(LoadedDocument(
                    content=text,
                    metadata={
                        # These metadata fields travel with chunks to ChromaDB
                        # When user gets answer, we show: "From gold_loan.pdf, Page 3"
                        "source": file_path,
                        "file_name": Path(file_path).name,
                        "file_type": "pdf",
                        "page": page_num + 1,           # 1-indexed (human readable)
                        "total_pages": total_pages,
                    },
                    source=file_path,
                    doc_type="pdf"
                ))

        print(f"   ✅ Loaded {len(documents)} non-empty pages")
        return documents

    # --------------------------------------------------------
    # DOCX LOADER
    # --------------------------------------------------------
    def _load_docx(self, file_path: str) -> List[LoadedDocument]:
        """
        WHAT: Reads Word document (.docx), extracts all paragraph text.

        WHY single document (not per page):
          Word docs don't have fixed pages like PDFs.
          Content flows based on font/zoom.
          So we treat whole doc as one unit.

        HOW:
          python-docx reads each paragraph
          Filter empty paragraphs
          Join all into single text string
          Return as one LoadedDocument

        Args:
            file_path: Full path to .docx file

        Returns:
            List with ONE LoadedDocument (whole document)
        """
        print(f"📝 Loading DOCX: {Path(file_path).name}")

        doc = DocxDocument(file_path)

        # Extract text from each paragraph, skip empty ones
        paragraphs = [
            para.text
            for para in doc.paragraphs
            if para.text.strip()    # Skip empty paragraphs
        ]

        # Join all paragraphs with newlines
        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            raise ValueError(f"❌ DOCX file appears to be empty: {file_path}")

        print(f"   ✅ Loaded {len(paragraphs)} paragraphs")

        return [LoadedDocument(
            content=full_text,
            metadata={
                "source": file_path,
                "file_name": Path(file_path).name,
                "file_type": "docx",
                "paragraph_count": len(paragraphs),
            },
            source=file_path,
            doc_type="docx"
        )]

    # --------------------------------------------------------
    # URL LOADER
    # --------------------------------------------------------
    def _load_url(self, url: str) -> List[LoadedDocument]:
        """
        WHAT: Fetches webpage, extracts readable text content.

        WHY:  Portfolio project use case:
              Load your company FAQ page, documentation site,
              Wikipedia articles, news articles etc.
              No need to download manually — just pass URL!

        HOW:
          httpx.get(url) → fetch raw HTML
          BeautifulSoup parses HTML
          Remove noise: scripts, styles, nav, footer
          Extract clean readable text
          Return as LoadedDocument

        Args:
            url: Full URL starting with http:// or https://

        Returns:
            List with ONE LoadedDocument (full page content)
        """
        print(f"🌐 Fetching URL: {url}")

        # Fetch the webpage
        # timeout=30 → don't wait more than 30 seconds
        # follow_redirects=True → handle http→https redirects
        response = httpx.get(
            url,
            timeout=30,
            follow_redirects=True,
            headers={
                # Pretend to be a browser
                # Some sites block requests without User-Agent
                "User-Agent": "Mozilla/5.0 (compatible; RAGBot/1.0)"
            }
        )

        # Raise error if page not found (404) or server error (500)
        response.raise_for_status()

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(response.text, "html.parser")

        # Remove elements that are noise (not actual content)
        # Scripts → JavaScript code (not useful text)
        # Styles  → CSS code (not useful text)
        # Nav     → Navigation menu links
        # Footer  → Copyright, links (not useful text)
        for noise_element in soup(["script", "style", "nav",
                                   "footer", "header", "aside"]):
            noise_element.decompose()   # Remove from HTML tree

        # Extract clean text
        # separator="\n" → add newline between elements
        # strip=True     → remove extra whitespace
        text = soup.get_text(separator="\n", strip=True)

        # Get page title for metadata
        title_tag = soup.find("title")
        page_title = title_tag.get_text() if title_tag else url

        print(f"   ✅ Loaded URL: '{page_title}' ({len(text)} chars)")

        return [LoadedDocument(
            content=text,
            metadata={
                "source": url,
                "file_name": page_title,
                "file_type": "url",
                "url": url,
                "title": page_title,
            },
            source=url,
            doc_type="url"
        )]

    # --------------------------------------------------------
    # TXT LOADER
    # --------------------------------------------------------
    def _load_txt(self, file_path: str) -> List[LoadedDocument]:
        """
        WHAT: Reads plain text file (.txt).

        WHY:  Simplest format — no parsing needed.
              Good for: notes, logs, raw data files.

        HOW:
          open() with utf-8 encoding
          Read entire content
          Return as one LoadedDocument

        Args:
            file_path: Full path to .txt file

        Returns:
            List with ONE LoadedDocument
        """
        print(f"📃 Loading TXT: {Path(file_path).name}")

        # encoding="utf-8" → handles Tamil, Hindi, special chars
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        if not content.strip():
            raise ValueError(f"❌ Text file is empty: {file_path}")

        print(f"   ✅ Loaded {len(content)} characters")

        return [LoadedDocument(
            content=content,
            metadata={
                "source": file_path,
                "file_name": Path(file_path).name,
                "file_type": "txt",
                "char_count": len(content),
            },
            source=file_path,
            doc_type="txt"
        )]




















































































# import os
# from pathlib import Path
# from typing import List
# from dataclasses import dataclass, field

# import pypdf
# from docx import Document as DocxDocument

# import httpx
# from bs4 import BeautifulSoup

# from src.config import Settings

# @dataclass
# class loadeddocument:
#     content: str
#     metadata: dict
#     source: str
#     doc_type: str

# class documentloader:

#     def load(self,source:str) ->List[loadeddocument]:
#         source=source.strip()

#         if source.startswith("http://") or source.startswith("https://"):
#             return self._load_url(source)
        
#         path=Path(source)

#         if not path.exists():
#             raise FileNotFoundError(
#                 f' file not found : {source}\n'
#                 f' check if path is correct'
#             )
#         extention=path.suffix.lower()

#         if extention=='.pdf':
#             return self._load_pdf(source)
#         elif extention in ['.docx','.doc']:
#             return self._load_docx(source)
#         elif extention =='.txt':
#             return self._load_txt(source)
#         else:
#             raise ValueError(
#                 f' unsupported file type : {extention}\n'
#                 f'supported type are : .pdf, .docx, .doc, .txt'
#             )
        

#     def _load_pdf(self,file_path:str)-> List[loadeddocument]:

#         document=[]

#         with open(file_path,'rb') as file:
#             pdf_reader=pypdf.PdfReader(file)
#             total_pages= len(pdf_reader.pages)

#             print(f' loading pdf : {Path(file_path).name}({total_pages} pages)')

#             for page_num , page in enumerate(pdf_reader.pages):

#                 text=page.extract_text()

#                 if not text or text.strip()=="":
#                     print(f' page {page_num +1} is empty , skipping')
#                     continue

#                 documents.append(loadeddocument(
#                     content=text,
#                     metadata={
#                             "source": file_path,
#                             "file_name": Path(file_path).name,
#                             "file_type": "pdf",
#                             "page": page_num +1,
#                             "total_pages": total_pages,

#                         },
#                         source=file_path,
#                         doc_type="pdf"
#                     ))   
#             print(f' loaded {len(documents)} non empty pages')
#             return documents
        
#         def _load_docx(self,file_path: str) -> List[loadeddocument]:
#             print(f' ;loading deocx: {Path(file_path).name}')
#             doc=DocxDocument(file_path)
#             paragraphs=[
#                 para.text
#                 for para in doc.paragraphs
#                 if para.text.strip()
#             ]

#             full_text="\n\n".join(paragraphs)

#             if not full_text.strip():
#                 raise ValueError(f' docx file appears to be empty : {file_path}')
            
#             print(f' loaded {len(paragraphs)}pararaphs')
#             return [loadeddoument(
#                 content=full_text,
#                 metadata={
#                     "source": file_path,
#                     "file_name": Path(file_path).name,
#                     "file_type": "docx",
#                     "paragraph_count": len(paaragraphs)
#                 },
#                 source=file_path,
#                 doc_type='docx'
#             )]
#         def _load_url(self,url:str)-> List[loadeddocument]:
#              print(f' fetching url : {url}')

#              response=httpx.get(
#                  url,timeout=30,
#                  follow_redirects=True,
#                  headers={
#                      "User-Agent": "Mozilla/5.0 (compatible; RAGBot/1.0)"

#                  }
#              )
#              response.raise_for_status()

#              soup=BeautifulSoup(response.text, "html.parser")

#              for noise_element in soup(["script","style",'nav','footer','header','aside']):
#                  noise_element.decompose()
#              text=soup.get_text(separator="\n", strip=True)
#              title_tag = soup.find("title")
#              page_title= title_tag.get_text() if title_tag else url

#              print(f' loaded url {page_title} ({len(text)} characters)')
#              return [loadeddocument(
#                  content=text,
#                  metadata={
#                      "source": url,
#                      "file_name": page_title,
#                      "file_type": "url",
#                      "url": url,
#                      "title": page_title,

#                  },source=url,
#                  doc_type="url"
#              )]
        
#         def _load_txt(self,file_path: str) -> List[loadeddocument]:
#             print(f' loading text file {Path(file_path).name}')
        
#         with open (file_path, "r", encoding ="utf-8") as f:
#             content = f.read()

#             if not content.strip():
#                 raise ValueError(f' text file is empty : {file_path}')
            
#             print(f' ;loaded {len(content)} characters')

#             return [loadeddocument(
#                 content=content,
#                 metadata={
#                     "source": file_path,
#                     "file_name": Path(file_path).name,
#                     "file_type": "txt",
#                     "character_count": len(content),
#                 },
#                 source=file_path,
#                 doc_type="txt"
#             )]